"""
Serviço para gerar documentos Word a partir de modelos.
"""
import os
import json
from docx import Document
from typing import Dict, Any, List
import re


def substituir_texto_em_paragrafo(paragraph, tag, valor):
    """
    Substitui uma tag por um valor em um parágrafo, preservando formatação completa.
    """
    if not paragraph:
        return False
    
    # Junta todo o texto do parágrafo de todos os runs
    texto_completo = ''.join([run.text for run in paragraph.runs])
    
    if tag not in texto_completo:
        return False
    
    # Substitui a tag pelo valor no texto completo
    novo_texto = texto_completo.replace(tag, str(valor))
    
    # Preserva a formatação: mantém o primeiro run e remove os outros
    if paragraph.runs:
        primeiro_run = paragraph.runs[0]
        
        # Preserva propriedades da fonte do primeiro run
        fonte_original = primeiro_run.font.name
        tamanho_original = primeiro_run.font.size
        negrito_original = primeiro_run.bold
        italico_original = primeiro_run.italic
        sublinhado_original = primeiro_run.underline
        
        # Remove runs extras (do último para o primeiro, exceto o primeiro)
        runs_para_remover = list(paragraph.runs[1:])
        for run in runs_para_remover:
            p_element = paragraph._element
            run_element = run._element
            p_element.remove(run_element)
        
        # Atualiza o texto do primeiro run
        primeiro_run.text = novo_texto
        
        # Restaura formatação (sem tentar definir color, pois não tem setter)
        if fonte_original:
            primeiro_run.font.name = fonte_original
        if tamanho_original:
            primeiro_run.font.size = tamanho_original
        primeiro_run.bold = negrito_original
        primeiro_run.italic = italico_original
        primeiro_run.underline = sublinhado_original
    else:
        # Se não houver runs, cria um novo
        paragraph.add_run(novo_texto)
    
    return True


def substituir_texto_em_documento(doc, tag, valor):
    """
    Substitui uma tag em todo o documento (parágrafos, tabelas, headers, footers).
    """
    substituido = False
    
    # Substitui em parágrafos principais
    for paragraph in doc.paragraphs:
        if substituir_texto_em_paragrafo(paragraph, tag, valor):
            substituido = True
    
    # Substitui em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if substituir_texto_em_paragrafo(paragraph, tag, valor):
                        substituido = True
    
    # Substitui em headers
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if substituir_texto_em_paragrafo(paragraph, tag, valor):
                substituido = True
        
        # Substitui em tabelas dos headers
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if substituir_texto_em_paragrafo(paragraph, tag, valor):
                            substituido = True
        
        # Substitui em footers
        footer = section.footer
        for paragraph in footer.paragraphs:
            if substituir_texto_em_paragrafo(paragraph, tag, valor):
                substituido = True
        
        # Substitui em tabelas dos footers
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if substituir_texto_em_paragrafo(paragraph, tag, valor):
                            substituido = True
    
    return substituido


def linha_contem_tag_sprint(row):
    """Verifica se uma linha contém tags de sprint."""
    tags_sprint = ['{SPRINT_OS}', '{OS_ID}', '{SPRINT_ID}', '{SPRINT_TIPO}', '{SPRINT_HST}', 
                   '{SPRINT_VALOR_H}', '{SPRINT_VALOR_TOTAL}', '{SPRINTS_HORAS}', '{SPRINTS_ HORAS}', '{ATIVIDADES}', '{ENTREGAVEIS}']
    texto_linha = ''
    for cell in row.cells:
        texto_linha += cell.text + ' '
    
    # IMPORTANTE: Verifica se há tags explícitas com chaves {}
    # Não aceita apenas palavras soltas como "Horas" - precisa ter a tag completa como "{SPRINTS_HORAS}"
    # Verifica se pelo menos uma tag completa está presente no texto
    tem_tag = False
    for tag in tags_sprint:
        if tag in texto_linha:
            tem_tag = True
            break
    
    # Verifica adicionalmente se há pelo menos uma chave {} no texto (garantia extra)
    tem_chaves = '{' in texto_linha and '}' in texto_linha
    
    # IMPORTANTE: Se a linha contém palavras de cabeçalho mas não tem tags, não é uma linha de template
    # Lista completa de palavras que aparecem em cabeçalhos de tabelas
    palavras_cabecalho = ['Fase', 'Sprint', 'Horas', 'OS*', 'Atividades', 'Entregáveis', 'Observação', 
                          'Perfil', 'Qtde', 'HSTs', 'Alocação', 'timebox', 'sprint (%)']
    # Verifica se a linha contém MÚLTIPLAS palavras de cabeçalho (indicando que é um cabeçalho real)
    palavras_encontradas = sum(1 for palavra in palavras_cabecalho if palavra in texto_linha)
    tem_apenas_cabecalho = palavras_encontradas >= 2 and not tem_chaves
    
    # CRITÉRIO FINAL: Só retorna True se tem tag E tem chaves E não é apenas cabeçalho
    resultado = tem_tag and tem_chaves and not tem_apenas_cabecalho
    
    return resultado


def linha_contem_tag_profissional(row):
    """Verifica se uma linha contém tags de profissional."""
    tags_prof = ['{PROF_TIPO}', '{PROF_QUANTIDADE}', '{PROF_HORAS}', '{PROF_QTD}', '{PORCENTAGEM}']
    texto_linha = ''
    for cell in row.cells:
        texto_linha += cell.text + ' '
    return any(tag in texto_linha for tag in tags_prof)


def linha_contem_tag_numerada(row, sprint_num):
    """Verifica se uma linha contém tags numeradas de uma sprint específica."""
    import re
    texto_linha = ''
    for cell in row.cells:
        texto_linha += cell.text + ' '
    # Procura por tags numeradas como {SPRINT_ID_1}, {PROF_TIPO_1_1}, etc.
    pattern = r'\{[A-Z_]+_' + str(sprint_num) + r'(_\d+)?\}'
    return bool(re.search(pattern, texto_linha))


def identificar_sprint_num_na_linha(row):
    """Identifica o número da sprint na linha baseado nas tags numeradas."""
    import re
    texto_linha = ''
    for cell in row.cells:
        texto_linha += cell.text + ' '
    # Procura por padrão {SPRINT_ID_N} ou {SPRINT_TIPO_N}
    match = re.search(r'\{SPRINT_ID_(\d+)\}', texto_linha)
    if match:
        return int(match.group(1))
    match = re.search(r'\{SPRINT_TIPO_(\d+)\}', texto_linha)
    if match:
        return int(match.group(1))
    return None


def identificar_prof_num_na_linha(row, sprint_num):
    """Identifica o número do profissional na linha baseado nas tags numeradas."""
    import re
    texto_linha = ''
    for cell in row.cells:
        texto_linha += cell.text + ' '
    # Procura por padrão {PROF_TIPO_N_M}
    pattern = r'\{PROF_TIPO_' + str(sprint_num) + r'_(\d+)\}'
    match = re.search(pattern, texto_linha)
    if match:
        return int(match.group(1))
    return None


def duplicar_linha_tabela(table, linha_template_index):
    """Duplica uma linha de tabela mantendo formatação completa."""
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from copy import deepcopy
    
    linha_template = table.rows[linha_template_index]
    
    # IMPORTANTE: add_row() cria uma linha baseada na última linha da tabela
    # Para garantir que temos o número correto de células, vamos usar a linha template como referência
    # Primeiro, adiciona uma linha baseada na linha template (não na última linha)
    nova_linha = table.add_row()
    
    # Valida que o número de células corresponde
    num_celulas_template = len(linha_template.cells)
    num_celulas_nova = len(nova_linha.cells)
    
    # Se o número de células não corresponde, ajusta removendo células extras
    if num_celulas_nova != num_celulas_template:
        print(f"[DEBUG] Aviso: Número de células diferente! Template: {num_celulas_template}, Nova: {num_celulas_nova}")
        # Se a nova linha tem mais células, remove as extras (da direita para esquerda)
        if num_celulas_nova > num_celulas_template:
            # Remove células extras do final - precisa fazer em ordem reversa para não quebrar índices
            celulas_para_remover = num_celulas_nova - num_celulas_template
            # Pega a lista de células antes de remover
            celulas_para_remover_lista = list(nova_linha.cells)[num_celulas_template:]
            # Remove cada célula extra
            for cell_extra in celulas_para_remover_lista:
                try:
                    nova_linha._element.remove(cell_extra._element)
                    print(f"[DEBUG] Removida célula extra")
                except Exception as e:
                    print(f"[DEBUG] Erro ao remover célula extra: {e}")
        # Se a nova linha tem menos células, isso não deveria acontecer com add_row()
        elif num_celulas_nova < num_celulas_template:
            print(f"[DEBUG] Erro: Nova linha tem menos células que o template! Isso não deveria acontecer.")
    
    # Agora valida novamente após remoção
    num_celulas_nova_final = len(nova_linha.cells)
    if num_celulas_nova_final != num_celulas_template:
        print(f"[DEBUG] Erro crítico: Após ajuste, ainda há diferença! Template: {num_celulas_template}, Nova: {num_celulas_nova_final}")
        # Se ainda não corresponde, processa apenas até o mínimo
        num_celulas_processar = min(num_celulas_template, num_celulas_nova_final)
    else:
        num_celulas_processar = num_celulas_template
    
    # Copia conteúdo e formatação de cada célula
    for i in range(num_celulas_processar):
        cell_template = linha_template.cells[i]
        nova_cell = nova_linha.cells[i]
        is_ultima_coluna = (i == num_celulas_processar - 1)
        
        # Copia propriedades da célula (cor de fundo, etc.)
        # Copia shading (cor de fundo) - método mais robusto
        if cell_template._element.tcPr is not None:
            shd_element = cell_template._element.tcPr.find(qn('w:shd'))
            if shd_element is not None:
                if nova_cell._element.tcPr is None:
                    nova_cell._element.add_tcPr()
                # Remove shading existente se houver
                shd_existente = nova_cell._element.tcPr.find(qn('w:shd'))
                if shd_existente is not None:
                    nova_cell._element.tcPr.remove(shd_existente)
                # Copia o shading
                nova_cell._element.tcPr.append(deepcopy(shd_element))
        
        # Limpa a nova célula completamente (remove todos os parágrafos exceto o primeiro)
        # Remove parágrafos extras
        while len(nova_cell.paragraphs) > 1:
            nova_cell._element.remove(nova_cell.paragraphs[-1]._element)
        
        # Limpa o primeiro parágrafo mas mantém sua estrutura
        primeiro_para_novo = nova_cell.paragraphs[0]
        for run in primeiro_para_novo.runs:
            primeiro_para_novo._element.remove(run._element)
        primeiro_para_novo.text = ''
        
        # Para a última coluna (Observação), trata especialmente para evitar duplicação de "N/A"
        if is_ultima_coluna:
            # Para a última coluna, junta todo o texto primeiro para verificar se é "N/A"
            texto_completo_template = cell_template.text.strip()
            
            # Se o template tem "N/A", copia apenas uma vez preservando formatação
            if texto_completo_template.upper() in ['N/A', 'N / A', 'N/ A', 'N /A']:
                # Usa apenas o primeiro parágrafo e primeiro run como referência de formatação
                if cell_template.paragraphs and cell_template.paragraphs[0].runs:
                    run_referencia = cell_template.paragraphs[0].runs[0]
                    para_referencia = cell_template.paragraphs[0]
                    
                    # Limpa completamente
                    primeiro_para_novo.text = ''
                    for run in list(primeiro_para_novo.runs):
                        primeiro_para_novo._element.remove(run._element)
                    
                    # Adiciona apenas um run com "N/A"
                    novo_run = primeiro_para_novo.add_run('N/A')
                    
                    # Copia formatação do run de referência
                    if run_referencia._element.rPr is not None:
                        rPr_copy = deepcopy(run_referencia._element.rPr)
                        if novo_run._element.rPr is None:
                            novo_run._element.insert(0, rPr_copy)
                        else:
                            novo_run._element.remove(novo_run._element.rPr)
                            novo_run._element.insert(0, rPr_copy)
                    else:
                        novo_run.bold = run_referencia.bold
                        novo_run.italic = run_referencia.italic
                        novo_run.underline = run_referencia.underline
                        if run_referencia.font.name:
                            novo_run.font.name = run_referencia.font.name
                        if run_referencia.font.size:
                            novo_run.font.size = run_referencia.font.size
                    
                    # Copia alinhamento do parágrafo de referência
                    if para_referencia._element.pPr is not None:
                        pPr_copy = deepcopy(para_referencia._element.pPr)
                        if primeiro_para_novo._element.pPr is None:
                            primeiro_para_novo._element.insert(0, pPr_copy)
                        else:
                            primeiro_para_novo._element.remove(primeiro_para_novo._element.pPr)
                            primeiro_para_novo._element.insert(0, pPr_copy)
                    elif para_referencia.alignment is not None:
                        primeiro_para_novo.alignment = para_referencia.alignment
                    
                    # Pula o resto do processamento para esta célula
                    continue
        
        # Para outras células, copia normalmente
        # Copia parágrafos preservando formatação completa
        for para_idx, para_template in enumerate(cell_template.paragraphs):
            if para_idx == 0:
                novo_para = nova_cell.paragraphs[0]
            else:
                novo_para = nova_cell.add_paragraph()
            
            # Copia propriedades do parágrafo (alinhamento) usando deepcopy do elemento pPr
            if para_template._element.pPr is not None:
                # Copia o elemento pPr completo para preservar todas as propriedades
                pPr_copy = deepcopy(para_template._element.pPr)
                if novo_para._element.pPr is None:
                    # Se não existe pPr, adiciona antes do elemento r
                    novo_para._element.insert(0, pPr_copy)
                else:
                    # Substitui o pPr existente
                    novo_para._element.remove(novo_para._element.pPr)
                    novo_para._element.insert(0, pPr_copy)
            elif para_template.alignment is not None:
                # Fallback: se não tem pPr mas tem alinhamento, define apenas o alinhamento
                novo_para.alignment = para_template.alignment
            
            # Limpa runs existentes do novo parágrafo (se houver)
            for run in list(novo_para.runs):
                novo_para._element.remove(run._element)
            
            # Copia runs preservando formatação completa
            for run_template in para_template.runs:
                texto_run = run_template.text
                novo_run = novo_para.add_run(texto_run)
                
                # Copia formatação completa do run usando deepcopy do elemento rPr
                if run_template._element.rPr is not None:
                    rPr_copy = deepcopy(run_template._element.rPr)
                    if novo_run._element.rPr is None:
                        novo_run._element.insert(0, rPr_copy)
                    else:
                        novo_run._element.remove(novo_run._element.rPr)
                        novo_run._element.insert(0, rPr_copy)
                else:
                    # Fallback: copia propriedades básicas manualmente
                    novo_run.bold = run_template.bold
                    novo_run.italic = run_template.italic
                    novo_run.underline = run_template.underline
                    if run_template.font.name:
                        novo_run.font.name = run_template.font.name
                    if run_template.font.size:
                        novo_run.font.size = run_template.font.size
    
    # Retorna o índice da nova linha (última linha da tabela)
    return len(table.rows) - 1


def carregar_config_sprints():
    """
    Carrega o arquivo de configuração de sprints.
    Retorna um dicionário com as configurações ou um dicionário vazio se houver erro.
    """
    try:
        # Caminho do arquivo de configuração
        config_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'config', 'sprints_config.json')
        
        if not os.path.exists(config_path):
            print(f"[DEBUG] [WARN] Arquivo de configuração não encontrado: {config_path}")
            return {}
        
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
            print(f"[DEBUG] Configuração de sprints carregada com sucesso")
            return config
    except Exception as e:
        print(f"[DEBUG] [ERROR] Erro ao carregar configuração de sprints: {e}")
        return {}


def obter_atividades_por_tipo(tipo_sprint: str) -> str:
    """
    Obtém as atividades para um tipo de sprint específico.
    Retorna uma string vazia se o tipo não for encontrado.
    """
    config = carregar_config_sprints()
    tipos_sprint = config.get('tipos_sprint', {})
    
    # Normaliza o tipo (case-insensitive)
    tipo_normalizado = str(tipo_sprint).strip().lower()
    
    # Busca o tipo correspondente (case-insensitive)
    for tipo_config, dados in tipos_sprint.items():
        if tipo_config.lower() == tipo_normalizado:
            atividades = dados.get('atividades', '')
            print(f"[DEBUG] Atividades encontradas para tipo '{tipo_sprint}': {atividades}")
            return atividades
    
    print(f"[DEBUG] [WARN] Tipo de sprint '{tipo_sprint}' não encontrado no arquivo de configuração")
    return ''


def obter_entregaveis_por_tipo(tipo_sprint: str) -> str:
    """
    Obtém os entregáveis para um tipo de sprint específico.
    Retorna uma string vazia se o tipo não for encontrado.
    """
    config = carregar_config_sprints()
    tipos_sprint = config.get('tipos_sprint', {})
    
    # Normaliza o tipo (case-insensitive)
    tipo_normalizado = str(tipo_sprint).strip().lower()
    
    # Busca o tipo correspondente (case-insensitive)
    for tipo_config, dados in tipos_sprint.items():
        if tipo_config.lower() == tipo_normalizado:
            entregaveis = dados.get('entregaveis', '')
            print(f"[DEBUG] Entregáveis encontrados para tipo '{tipo_sprint}': {entregaveis}")
            return entregaveis
    
    print(f"[DEBUG] [WARN] Tipo de sprint '{tipo_sprint}' não encontrado no arquivo de configuração")
    return ''


def escrever_valor_em_celula(cell, valor):
    """Escreve um valor em uma célula preservando ao máximo a formatação original."""
    valor = '' if valor is None else str(valor)
    
    if not cell.paragraphs:
        cell.text = valor
        return
    
    paragrafo = cell.paragraphs[0]
    
    # Preserva formatação do primeiro run (se existir)
    fonte_ref = None
    tamanho_ref = None
    negrito_ref = False
    italico_ref = False
    rPr_ref = None
    
    if paragrafo.runs:
        primeiro_run = paragrafo.runs[0]
        fonte_ref = primeiro_run.font.name
        tamanho_ref = primeiro_run.font.size
        negrito_ref = primeiro_run.bold
        italico_ref = primeiro_run.italic
        if primeiro_run._element.rPr is not None:
            from copy import deepcopy
            rPr_ref = deepcopy(primeiro_run._element.rPr)
    
    # Remove todos os runs existentes
    for run in list(paragrafo.runs):
        paragrafo._element.remove(run._element)
    
    # Adiciona novo run com o valor
    novo_run = paragrafo.add_run(valor)
    
    # Aplica formatação preservada
    if rPr_ref is not None:
        if novo_run._element.rPr is None:
            novo_run._element.insert(0, rPr_ref)
        else:
            novo_run._element.remove(novo_run._element.rPr)
            novo_run._element.insert(0, rPr_ref)
    else:
        # Fallback: aplica propriedades básicas
        if fonte_ref:
            novo_run.font.name = fonte_ref
        if tamanho_ref:
            novo_run.font.size = tamanho_ref
        novo_run.bold = negrito_ref
        novo_run.italic = italico_ref


def preencher_linha_com_dados_sprint(row, sprint_data, tags_sprint):
    """Preenche uma linha com dados de uma sprint preservando formatação."""
    # Primeiro, substitui todas as tags preservando formatação
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for tag, campo in tags_sprint.items():
                # Lógica especial para diferenciar HST original de horas sprint do usuário
                if tag == '{SPRINT_HST}':
                    # Tabela 10: usa HST original do Redmine
                    valor = str(sprint_data.get('hst_redmine') or sprint_data.get('hst', ''))
                elif tag in ['{SPRINTS_HORAS}', '{SPRINTS_ HORAS}']:
                    # Tabela 6: usa horas_sprint (valor digitado pelo usuário)
                    valor = str(sprint_data.get('horas_sprint') or sprint_data.get('hst', ''))
                elif tag == '{ATIVIDADES}':
                    # Tabela 6: preenche atividades baseado no tipo da sprint (do arquivo de configuração)
                    tipo_sprint = sprint_data.get('tipo', '')
                    valor = obter_atividades_por_tipo(tipo_sprint)
                    print(f"[DEBUG] Preenchendo {tag} para sprint tipo '{tipo_sprint}': {valor}")
                elif tag == '{ENTREGAVEIS}':
                    # Tabela 6: preenche entregáveis baseado no tipo da sprint (do arquivo de configuração)
                    tipo_sprint = sprint_data.get('tipo', '')
                    valor = obter_entregaveis_por_tipo(tipo_sprint)
                    print(f"[DEBUG] Preenchendo {tag} para sprint tipo '{tipo_sprint}': {valor}")
                else:
                    valor = str(sprint_data.get(campo, ''))
                substituir_texto_em_paragrafo(paragraph, tag, valor)
    
    # IMPORTANTE: NÃO substitui valores diretos em células que não têm tags
    # Só preenche onde há tags explícitas para evitar sobrescrever dados normais
    # A substituição de tags já foi feita acima
    
    # SEMPRE normaliza a última coluna (Observação) para ter apenas "N/A" uma vez
    # Isso previne duplicação mesmo que tenha sido copiado incorretamente
    if len(row.cells) > 0:
        obs_cell = row.cells[-1]  # Última célula
        texto_obs = obs_cell.text.strip().upper()
        
        # Se a célula contém qualquer variação de "N/A", normaliza para apenas "N/A"
        if texto_obs and ('N' in texto_obs and '/' in texto_obs and 'A' in texto_obs):
            # Conta quantas vezes "N/A" aparece (pode estar duplicado)
            ocorrencias_na = texto_obs.count('N/A')
            ocorrencias_na_espacado = texto_obs.count('N / A') + texto_obs.count('N/ A') + texto_obs.count('N /A')
            
            # Se tem mais de uma ocorrência, está mal formatado, ou tem mais de 3 caracteres, normaliza
            if ocorrencias_na > 1 or ocorrencias_na_espacado > 0 or len(texto_obs.replace(' ', '').replace('/', '')) > 2:
                # Limpa completamente e adiciona apenas "N/A" preservando formatação
                if obs_cell.paragraphs:
                    primeiro_para = obs_cell.paragraphs[0]
                    
                    # Preserva alinhamento e formatação do parágrafo
                    alinhamento_original = primeiro_para.alignment
                    
                    # Preserva formatação do primeiro run existente (se houver)
                    fonte_ref = None
                    tamanho_ref = None
                    bold_ref = False
                    italic_ref = False
                    rPr_ref = None
                    
                    if primeiro_para.runs:
                        primeiro_run_original = primeiro_para.runs[0]
                        fonte_ref = primeiro_run_original.font.name
                        tamanho_ref = primeiro_run_original.font.size
                        bold_ref = primeiro_run_original.bold
                        italic_ref = primeiro_run_original.italic
                        if primeiro_run_original._element.rPr is not None:
                            from copy import deepcopy
                            rPr_ref = deepcopy(primeiro_run_original._element.rPr)
                    
                    # Se não tem run de referência, tenta usar outra célula da linha
                    if not rPr_ref and len(row.cells) > 0 and row.cells[0].paragraphs:
                        ref_para = row.cells[0].paragraphs[0]
                        if ref_para.runs:
                            ref_run = ref_para.runs[0]
                            fonte_ref = fonte_ref or ref_run.font.name
                            tamanho_ref = tamanho_ref or ref_run.font.size
                            bold_ref = bold_ref or ref_run.bold
                            italic_ref = italic_ref or ref_run.italic
                            if ref_run._element.rPr is not None:
                                from copy import deepcopy
                                rPr_ref = deepcopy(ref_run._element.rPr)
                    
                    # Remove todos os runs
                    for run in list(primeiro_para.runs):
                        primeiro_para._element.remove(run._element)
                    
                    # Adiciona novo run com "N/A"
                    novo_run = primeiro_para.add_run('N/A')
                    
                    # Aplica formatação preservada
                    if rPr_ref:
                        if novo_run._element.rPr is None:
                            novo_run._element.insert(0, rPr_ref)
                        else:
                            novo_run._element.remove(novo_run._element.rPr)
                            novo_run._element.insert(0, rPr_ref)
                    else:
                        # Fallback: copia propriedades básicas
                        if fonte_ref:
                            novo_run.font.name = fonte_ref
                        if tamanho_ref:
                            novo_run.font.size = tamanho_ref
                        novo_run.bold = bold_ref
                        novo_run.italic = italic_ref
                    
                    # Restaura alinhamento
                    if alinhamento_original is not None:
                        primeiro_para.alignment = alinhamento_original


def preencher_linha_com_dados_profissional(row, prof_data, tags_prof):
    """Preenche uma linha com dados de um profissional."""
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for tag, campo in tags_prof.items():
                valor = str(prof_data.get(campo, ''))
                substituir_texto_em_paragrafo(paragraph, tag, valor)


def celula_contem_tag(cell, tag):
    """
    Verifica se uma célula contém uma tag específica.
    Retorna True se a tag estiver presente no texto da célula.
    """
    texto_completo = ''.join([
        run.text for paragraph in cell.paragraphs 
        for run in paragraph.runs
    ])
    return tag in texto_completo


def preencher_linha_item7(row, sprint_data, prof_data, tags_sprint, tags_prof, mostrar_sprint=True):
    """
    Preenche uma linha da Tabela 7 com dados da sprint e do profissional.
    IMPORTANTE: Preenche APENAS onde há tags. Células sem tags (ex: "N/A") são mantidas como estão.
    """
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            # Substitui tags da sprint presentes na linha
            for tag, campo in tags_sprint.items():
                valor = str(sprint_data.get(campo, ''))
                substituir_texto_em_paragrafo(paragraph, tag, valor)
            
            # Substitui tags do profissional (se houver profissional associado)
            for tag, campo in tags_prof.items():
                valor = ''
                if prof_data:
                    if campo == 'percentual':
                        # Calcula a porcentagem automaticamente: (prof_horas / sprint_horas) * 100
                        # IMPORTANTE: Usa horas_sprint (horas digitadas pelo usuário) para a tabela 7, não HST original
                        try:
                            # Usa horas_sprint (horas da sprint digitadas pelo usuário), fallback para hst se não houver
                            sprint_horas_str = str(sprint_data.get('horas_sprint') or sprint_data.get('hst', '0')).strip()
                            prof_horas_str = str(prof_data.get('horas', '0')).strip()
                            
                            if sprint_horas_str and prof_horas_str:
                                sprint_horas = float(sprint_horas_str)
                                prof_horas = float(prof_horas_str)
                                
                                if sprint_horas > 0:
                                    porcentagem_valor = (prof_horas / sprint_horas) * 100
                                    # Formata como inteiro se for número inteiro, senão mantém 1 decimal, e adiciona %
                                    if porcentagem_valor == int(porcentagem_valor):
                                        valor = f"{int(porcentagem_valor)}%"
                                    else:
                                        valor = f"{porcentagem_valor:.1f}%"
                        except (ValueError, TypeError, ZeroDivisionError) as e:
                            print(f"[DEBUG] [WARN] Erro ao calcular porcentagem: {e}")
                            # Fallback: usa valor dos dados se disponível
                            valor = (
                                prof_data.get('percentual') or
                                prof_data.get('alocacao') or
                                prof_data.get('porcentagem') or
                                ''
                            )
                    else:
                        valor = prof_data.get(campo, '')
                substituir_texto_em_paragrafo(paragraph, tag, str(valor))
    
    # Atualiza as colunas de sprint apenas uma vez por grupo E APENAS se houver tags
    # IMPORTANTE: Preenche apenas onde há tags. Se a célula não tiver tag, mantém como está (ex: "N/A")
    if len(row.cells) >= 2 and mostrar_sprint:
        # Verifica se há tags de sprint nas células antes de preencher
        # Verifica todas as possíveis tags de sprint ID/OS
        tags_sprint_id_possiveis = ['{SPRINT_ID}', '{SPRINT_OS}', '{OS_ID}'] + [tag for tag in tags_sprint.keys() if 'SPRINT_ID' in tag or 'SPRINT_OS' in tag or 'OS_ID' in tag]
        tags_sprint_tipo_possiveis = ['{SPRINT_TIPO}'] + [tag for tag in tags_sprint.keys() if 'SPRINT_TIPO' in tag]
        
        tem_tag_sprint_id = any(celula_contem_tag(row.cells[0], tag) for tag in tags_sprint_id_possiveis)
        tem_tag_sprint_tipo = any(celula_contem_tag(row.cells[1], tag) for tag in tags_sprint_tipo_possiveis)
        
        if tem_tag_sprint_id:
            escrever_valor_em_celula(row.cells[0], sprint_data.get('sprint', ''))
        if tem_tag_sprint_tipo:
            escrever_valor_em_celula(row.cells[1], sprint_data.get('tipo', ''))


def preencher_tags_numeradas_item7(row, sprint_data, prof_data, sprint_num, prof_num, primeira_linha_grupo=False):
    """
    Preenche uma linha da Tabela 7 usando tags numeradas.
    Exemplo: {SPRINT_ID_1}, {SPRINT_TIPO_1}, {PROF_TIPO_1_1}, {PROF_QTD_1_1}, {PROF_HORAS_1_1}
    
    Args:
        primeira_linha_grupo: Se True, esta é a primeira linha do grupo de sprint (células mescladas)
    """
    import re
    from docx.oxml.ns import qn
    
    # Mapeamento de tags numeradas de sprint
    tags_sprint_numeradas = {
        f'{{SPRINT_ID_{sprint_num}}}': sprint_data.get('sprint', ''),
        f'{{SPRINT_TIPO_{sprint_num}}}': sprint_data.get('tipo', ''),
    }
    
    # Mapeamento de tags numeradas de profissional
    tags_prof_numeradas = {}
    if prof_data:
        # Trata variação {PROF_HORAS1_1} (sem underscore) e {PROF_HORAS_1_1} (com underscore)
        tags_prof_numeradas[f'{{PROF_TIPO_{sprint_num}_{prof_num}}}'] = prof_data.get('tipo', '')
        tags_prof_numeradas[f'{{PROF_QTD_{sprint_num}_{prof_num}}}'] = str(prof_data.get('quantidade', ''))
        tags_prof_numeradas[f'{{PROF_HORAS_{sprint_num}_{prof_num}}}'] = str(prof_data.get('horas', ''))
        tags_prof_numeradas[f'{{PROF_HORAS{sprint_num}_{prof_num}}}'] = str(prof_data.get('horas', ''))  # Variação sem underscore
        
        # Calcula a porcentagem automaticamente: (prof_horas / sprint_horas) * 100
        # IMPORTANTE: Usa horas_sprint (horas digitadas pelo usuário) para a tabela 7, não HST original
        porcentagem_calculada = ''
        try:
            # Usa horas_sprint (horas da sprint digitadas pelo usuário), fallback para hst se não houver
            sprint_horas_str = str(sprint_data.get('horas_sprint') or sprint_data.get('hst', '0')).strip()
            prof_horas_str = str(prof_data.get('horas', '0')).strip()
            
            print(f"[DEBUG] Calculando porcentagem - Sprint {sprint_num}, Prof {prof_num}: sprint_horas={sprint_horas_str}, prof_horas={prof_horas_str}")
            
            if sprint_horas_str and prof_horas_str:
                sprint_horas = float(sprint_horas_str)
                prof_horas = float(prof_horas_str)
                
                if sprint_horas > 0:
                    porcentagem_valor = (prof_horas / sprint_horas) * 100
                    # Formata como inteiro se for número inteiro, senão mantém 1 decimal, e adiciona %
                    if porcentagem_valor == int(porcentagem_valor):
                        porcentagem_calculada = f"{int(porcentagem_valor)}%"
                    else:
                        porcentagem_calculada = f"{porcentagem_valor:.1f}%"
                    print(f"[DEBUG] Porcentagem calculada: {porcentagem_calculada}")
        except (ValueError, TypeError, ZeroDivisionError) as e:
            print(f"[DEBUG] [WARN] Erro ao calcular porcentagem: {e}")
            # Fallback: usa valor dos dados se disponível
            porcentagem_calculada = (
                prof_data.get('percentual') or
                prof_data.get('alocacao') or
                prof_data.get('porcentagem') or
                ''
            )
        
        # Adiciona tanto a tag genérica quanto a tag numerada
        tags_prof_numeradas['{PORCENTAGEM}'] = porcentagem_calculada
        tag_porcentagem_numerada = f'{{PORCENTAGEM_{sprint_num}_{prof_num}}}'
        tags_prof_numeradas[tag_porcentagem_numerada] = porcentagem_calculada
        print(f"[DEBUG] Tag de porcentagem adicionada: {tag_porcentagem_numerada} = {porcentagem_calculada}")
    else:
        # Se não há profissional, deixa tags vazias (serão ignoradas na substituição)
        print(f"[DEBUG] Sem dados de profissional para sprint {sprint_num}, preenchendo apenas tags de sprint")
    
    # Preenche células de sprint primeiro (apenas na primeira linha do grupo).
    # IMPORTANTE: Preenche APENAS onde há tags. Se a célula não tiver tag, mantém como está (ex: "N/A")
    # Importante: NÃO limpar as células nas linhas seguintes, pois em cenários
    # com mesclagem vertical o Word reutiliza o mesmo elemento de célula
    # para todas as linhas mescladas. Se limparmos nas linhas \"de baixo\",
    # apagamos também o conteúdo da primeira linha.
    if primeira_linha_grupo and len(row.cells) >= 2:
        # Verifica se há tags de sprint nas células antes de preencher
        # IMPORTANTE: Preenche APENAS onde há tags. Se não houver tag, mantém como está (ex: "N/A")
        tag_sprint_id = f'{{SPRINT_ID_{sprint_num}}}'
        tag_sprint_tipo = f'{{SPRINT_TIPO_{sprint_num}}}'
        
        # Também verifica tags genéricas caso existam
        tags_sprint_id_possiveis = [
            tag_sprint_id,
            '{SPRINT_ID}',
            '{SPRINT_OS}',
            '{OS_ID}'
        ]
        tags_sprint_tipo_possiveis = [
            tag_sprint_tipo,
            '{SPRINT_TIPO}'
        ]
        
        # Preenche apenas se a célula contém alguma tag correspondente
        tem_tag_id = any(celula_contem_tag(row.cells[0], tag) for tag in tags_sprint_id_possiveis)
        tem_tag_tipo = any(celula_contem_tag(row.cells[1], tag) for tag in tags_sprint_tipo_possiveis)
        
        if tem_tag_id:
            valor_sprint = str(sprint_data.get('sprint', ''))
            escrever_valor_em_celula(row.cells[0], valor_sprint)
            print(f"[DEBUG] Preenchida célula 0 com sprint {valor_sprint} (tag encontrada)")
        else:
            print(f"[DEBUG] Célula 0 não contém tag de sprint, mantendo conteúdo original (ex: 'N/A')")
        
        if tem_tag_tipo:
            valor_tipo = str(sprint_data.get('tipo', ''))
            escrever_valor_em_celula(row.cells[1], valor_tipo)
            print(f"[DEBUG] Preenchida célula 1 com tipo {valor_tipo} (tag encontrada)")
        else:
            print(f"[DEBUG] Célula 1 não contém tag de tipo, mantendo conteúdo original (ex: 'N/A')")
    
    # Substitui todas as tags numeradas na linha
    # IMPORTANTE: Tags de sprint nas células 0 e 1 já foram preenchidas diretamente acima
    # Mas ainda precisa substituir tags de sprint em outras células (caso existam)
    # IMPORTANTE: Também substitui tags nas células 0 e 1 caso não tenham sido preenchidas diretamente
    for cell_idx, cell in enumerate(row.cells):
        for paragraph in cell.paragraphs:
            # Substitui tags de sprint numeradas em TODAS as células (incluindo 0 e 1)
            # Isso garante que tags como {SPRINT_ID_1} sejam substituídas mesmo se não foram preenchidas diretamente
            for tag, valor in tags_sprint_numeradas.items():
                if substituir_texto_em_paragrafo(paragraph, tag, str(valor)):
                    print(f"[DEBUG] Tag de sprint {tag} substituída na célula {cell_idx} com valor: {valor}")
            
            # IMPORTANTE: Também substitui tags genéricas de sprint caso existam
            # Isso garante compatibilidade com templates que usam tags genéricas
            if cell_idx == 0:
                # Célula 0: substitui tags de sprint ID
                for tag_gen in ['{SPRINT_ID}', '{SPRINT_OS}', '{OS_ID}']:
                    if tag_gen in paragraph.text:
                        valor_sprint = str(sprint_data.get('sprint', ''))
                        if substituir_texto_em_paragrafo(paragraph, tag_gen, valor_sprint):
                            print(f"[DEBUG] Tag genérica de sprint {tag_gen} substituída na célula 0 com valor: {valor_sprint}")
            elif cell_idx == 1:
                # Célula 1: substitui tags de tipo de sprint
                if '{SPRINT_TIPO}' in paragraph.text:
                    valor_tipo = str(sprint_data.get('tipo', ''))
                    if substituir_texto_em_paragrafo(paragraph, '{SPRINT_TIPO}', valor_tipo):
                        print(f"[DEBUG] Tag genérica de tipo {{SPRINT_TIPO}} substituída na célula 1 com valor: {valor_tipo}")
            
            # Substitui tags de profissional numeradas (apenas se houver profissional)
            if tags_prof_numeradas:
                # IMPORTANTE: Porcentagem pode estar em qualquer célula, então processa todas as células
                for tag, valor in tags_prof_numeradas.items():
                    # Se for tag de porcentagem, substitui em qualquer célula
                    if '{PORCENTAGEM' in tag:
                        if substituir_texto_em_paragrafo(paragraph, tag, str(valor)):
                            print(f"[DEBUG] Tag {tag} substituída na célula {cell_idx} com valor: {valor}")
                    # Para outras tags de profissional, só substitui nas células >= 2 (exceto sprint)
                    elif cell_idx >= 2:
                        if substituir_texto_em_paragrafo(paragraph, tag, str(valor)):
                            print(f"[DEBUG] Tag de profissional {tag} substituída na célula {cell_idx} com valor: {valor}")


def preencher_plano_trabalho(
    modelo_path: str,
    dados_demanda: Dict[str, Any],
    dados_sprints: List[Dict[str, Any]],
    dados_profissionais: Dict[str, List[Dict[str, Any]]],
    dados_projeto: Dict[str, Any] = None
) -> Document:
    """
    Preenche o modelo de Plano de Trabalho com os dados fornecidos.
    
    Args:
        modelo_path: Caminho para o arquivo modelo .docx
        dados_demanda: Dicionário com dados da demanda (demanda, pt, nome, valor_demanda)
        dados_sprints: Lista de dicionários com dados das sprints
        dados_profissionais: Dicionário onde a chave é o ID da sprint e o valor é lista de profissionais
        dados_projeto: Dicionário com dados do projeto (gestor, gerente, introdução, etc.)
        
    Returns:
        Documento Word preenchido
    """
    if dados_projeto is None:
        dados_projeto = {}
    # Abre o documento modelo
    doc = Document(modelo_path)
    
    # Log para debug
    print(f"[DEBUG] Processando documento...")
    print(f"[DEBUG] Dados demanda: {dados_demanda}")
    print(f"[DEBUG] Dados sprints: {dados_sprints}")
    print(f"[DEBUG] Dados profissionais: {dados_profissionais}")
    
    # Função auxiliar para listar todas as tags encontradas no documento
    def listar_tags_no_documento(doc):
        """Lista todas as tags encontradas no documento para debug."""
        tags_encontradas = set()
        texto_completo = ''
        
        # Parágrafos principais
        for p in doc.paragraphs:
            texto_completo += p.text + ' '
        
        # Tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto_completo += cell.text + ' '
        
        # Headers e footers
        for section in doc.sections:
            for p in section.header.paragraphs:
                texto_completo += p.text + ' '
            for p in section.footer.paragraphs:
                texto_completo += p.text + ' '
        
        # Procura por padrão de tags {TAG}
        import re
        padrao = r'\{[A-Z_0-9]+\}'
        tags_encontradas = set(re.findall(padrao, texto_completo))
        
        print(f"[DEBUG] Tags encontradas no documento: {sorted(tags_encontradas)}")
        return tags_encontradas
    
    # Lista tags encontradas
    tags_encontradas = listar_tags_no_documento(doc)
    
    # Mapeamento de tags simples para valores
    tags_simples = {
        '{DEMANDA}': str(dados_demanda.get('demanda', '')),
        '{PT}': str(dados_demanda.get('pt', '')),
        '{NOME_PROJETO}': str(dados_demanda.get('nome', '')),
        '{VALOR_DEMANDA}': str(dados_demanda.get('valor_demanda', '')),
    }

    # Tags do projeto (Gestor e Gerente)
    if dados_projeto:
        gestor_nome = str(dados_projeto.get('gestorNome', ''))
        gestor_email = str(dados_projeto.get('gestorEmail', ''))
        gestor_celular = str(dados_projeto.get('gestorCelular', ''))
        gerente_nome = str(dados_projeto.get('gerenteNome', ''))
        gerente_email = str(dados_projeto.get('gerenteEmail', ''))
        gerente_telefone = str(dados_projeto.get('gerenteTelefone', ''))
        descricao_projeto = str(dados_projeto.get('introducaoProjeto', ''))
        
        tags_simples.update({
            '{GESTOR}': gestor_nome,
            '{GESTOR_EMAIL}': gestor_email,
            '{GESTOR_CELULAR}': gestor_celular,
            '{GESTOR_CELULAR}}': gestor_celular,  # Trata erro de dupla chave no template
            '{GERENTE}': gerente_nome,
            '{GERENTE_EMAIL}': gerente_email,
            '{GERENTE_CELULAR}': gerente_telefone,
            '{GERENTE_CELULAR}}': gerente_telefone,  # Trata erro de dupla chave no template
            '{DESCRICAO_PROJETO}': descricao_projeto,
        })
        print(f"[DEBUG] Tags do projeto adicionadas:")
        print(f"  GESTOR: {gestor_nome}")
        print(f"  GERENTE: {gerente_nome}")
        print(f"  DESCRICAO_PROJETO: {len(descricao_projeto)} caracteres")
    else:
        print(f"[DEBUG] [WARN] Dados do projeto não fornecidos - tags de projeto não serão preenchidas")

    # Tags especiais independentes dos dados da demanda
    # Ex.: {{data}} e {DATA} no histórico de revisões (data de geração do documento)
    from datetime import datetime
    data_hoje = datetime.today().strftime('%d/%m/%Y')
    tags_simples.update({
        '{{data}}': data_hoje,
        '{DATA}': data_hoje,  # Tag alternativa em maiúsculas
        '{data}': data_hoje,  # Tag alternativa em minúsculas
    })
    print(f"[DEBUG] Data de hoje definida: {data_hoje}")
    
    # Calcula o total de HSTs de todas as sprints (para Item 10)
    # IMPORTANTE: Usa o HST original do Redmine (hst_redmine se existir, senão hst)
    total_hst = 0
    if dados_sprints:
        for sprint in dados_sprints:
            # Prioriza hst_redmine (original do Redmine) sobre hst
            hst_str = str(sprint.get('hst_redmine') or sprint.get('hst', '0')).strip()
            if hst_str:
                try:
                    # Tenta converter para float (pode ser "160" ou "160.0")
                    hst_valor = float(hst_str)
                    total_hst += hst_valor
                except (ValueError, TypeError):
                    print(f"[DEBUG] [WARN] HST inválido na sprint {sprint.get('sprint', 'N/A')}: '{hst_str}'")
    tags_simples.update({
        '{TOTAL_HST}': str(int(total_hst)) if total_hst == int(total_hst) else str(total_hst),
    })
    print(f"[DEBUG] Total de HSTs calculado: {total_hst}")
    
    # Log das tags que serão substituídas
    print(f"[DEBUG] Tags simples a substituir:")
    for tag, valor in tags_simples.items():
        print(f"  {tag} -> '{valor}'")
    
    # Substitui tags simples em todo o documento
    for tag, valor in tags_simples.items():
        substituido = substituir_texto_em_documento(doc, tag, valor)
        if substituido:
            print(f"[DEBUG] [OK] Tag {tag} substituída com sucesso")
        else:
            print(f"[DEBUG] [WARN] Tag {tag} não encontrada no documento")
    
    # -----------------------------
    # 1) TAGS SIMPLES E SPRINTS
    # -----------------------------
    # Processa sprints - mapeamento de tags
    tags_sprint = {
        '{SPRINT_OS}': 'os',
        '{OS_ID}': 'os',  # Tag alternativa para OS
        '{SPRINT_ID}': 'sprint',
        '{SPRINT_TIPO}': 'tipo',
        '{SPRINT_HST}': 'hst',
        '{SPRINT_VALOR_H}': 'valor_h_sprint',
        '{SPRINT_VALOR_TOTAL}': 'valor_total',
        '{SPRINTS_HORAS}': 'hst',  # Tag alternativa para horas
        '{SPRINTS_ HORAS}': 'hst',
        '{ATIVIDADES}': 'atividades',  # Tag para atividades (preenchida do arquivo de configuração)
        '{ENTREGAVEIS}': 'entregaveis',  # Tag para entregáveis (preenchida do arquivo de configuração)
    }
    
    # Processa sprints em tabelas
    if dados_sprints:
        print(f"[DEBUG] Processando {len(dados_sprints)} sprint(s)...")
        
        # Para cada tabela no documento
        for table_idx, table in enumerate(doc.tables):
            # Encontra linhas que contêm tags de sprint
            # IMPORTANTE: Pula o cabeçalho (linha 0) e procura tags nas linhas seguintes
            linhas_template_sprint = []
            for row_idx, row in enumerate(table.rows):
                # Pula o cabeçalho (linha 0)
                if row_idx == 0:
                    continue
                
                # IMPORTANTE: Verifica se é uma linha de cabeçalho (contém múltiplas palavras de cabeçalho sem tags)
                texto_linha_completo = ' '.join([cell.text.strip() for cell in row.cells])
                palavras_cabecalho = ['Fase', 'Sprint', 'Horas', 'OS*', 'Atividades', 'Entregáveis', 'Observação']
                palavras_encontradas = sum(1 for palavra in palavras_cabecalho if palavra in texto_linha_completo)
                tem_chaves = '{' in texto_linha_completo and '}' in texto_linha_completo
                
                # Se tem múltiplas palavras de cabeçalho mas não tem tags, é um cabeçalho - PULA
                if palavras_encontradas >= 2 and not tem_chaves:
                    print(f"[DEBUG] Tabela {table_idx}: Linha {row_idx} é cabeçalho (ignorada): {texto_linha_completo[:60]}")
                    continue
                
                # IMPORTANTE: Só adiciona se a linha CONTÉM tags de sprint
                # Não preenche linhas normais sem tags
                if linha_contem_tag_sprint(row) and not linha_contem_tag_profissional(row):
                    # Verifica se realmente tem tags com chaves {} (não apenas palavras soltas)
                    if '{' in texto_linha_completo and '}' in texto_linha_completo:
                        linhas_template_sprint.append(row_idx)
                        print(f"[DEBUG] Tabela {table_idx}: Linha {row_idx} contém tags de sprint: {row.cells[0].text[:50] if row.cells else 'N/A'}")
                    else:
                        print(f"[DEBUG] Tabela {table_idx}: Linha {row_idx} tem palavra relacionada mas não tem tags reais (ignorada): {texto_linha_completo[:50]}")
            
            # Se não encontrou linhas com tags, tenta identificar linhas COMPLETAMENTE VAZIAS
            # (útil quando o template tem linhas vazias sem tags, mas NÃO preenche linhas com dados normais)
            if not linhas_template_sprint and len(table.rows) > 1:
                print(f"[DEBUG] Tabela {table_idx}: Nenhuma tag encontrada, procurando linhas completamente vazias")
                # Pula o cabeçalho (primeira linha) e verifica se há linhas COMPLETAMENTE VAZIAS
                for row_idx in range(1, len(table.rows)):
                    row = table.rows[row_idx]
                    # Verifica se TODAS as células estão vazias (linha template)
                    todas_vazias = True
                    for cell in row.cells:
                        texto_cell = cell.text.strip()
                        # Se a célula tem conteúdo significativo (mais de 2 caracteres), não é template
                        if len(texto_cell) > 2:
                            todas_vazias = False
                            break
                    
                    if todas_vazias:
                        print(f"[DEBUG] Tabela {table_idx}: Linha {row_idx} está completamente vazia, adicionando tags")
                        linhas_template_sprint.append(row_idx)
                        # Adiciona tags temporárias para que o preenchimento funcione
                        if len(row.cells) > 0:
                            # Adiciona tag de sprint ID na primeira célula se estiver vazia
                            if not row.cells[0].text.strip():
                                row.cells[0].paragraphs[0].add_run('{SPRINT_ID}')
                            if len(row.cells) > 1 and not row.cells[1].text.strip():
                                row.cells[1].paragraphs[0].add_run('{SPRINT_TIPO}')
                            if len(row.cells) > 2:
                                row.cells[2].paragraphs[0].add_run('{SPRINTS_HORAS}')
                            if len(row.cells) > 3:
                                row.cells[3].paragraphs[0].add_run('{OS_ID}')
                            if len(row.cells) > 4:
                                row.cells[4].paragraphs[0].add_run('{ATIVIDADES}')
                            if len(row.cells) > 5:
                                row.cells[5].paragraphs[0].add_run('{ENTREGAVEIS}')
                        # Para após encontrar a primeira linha template vazia
                        break
            
            # Se encontrou linhas de template de sprint
            if linhas_template_sprint:
                print(f"[DEBUG] Tabela {table_idx}: Encontradas {len(linhas_template_sprint)} linha(s) de template de sprint")
                
                # Usa apenas as linhas necessárias (uma por sprint)
                num_sprints = len(dados_sprints)
                num_linhas_template = len(linhas_template_sprint)
                
                # Se há mais sprints que linhas template, cria linhas adicionais
                if num_sprints > num_linhas_template:
                    print(f"[DEBUG] Tabela {table_idx}: Criando {num_sprints - num_linhas_template} linha(s) adicional(is)")
                    # Usa a última linha template como modelo para criar novas linhas
                    ultima_linha_template_idx = linhas_template_sprint[-1]
                    ultima_linha_template = table.rows[ultima_linha_template_idx]
                    
                    for i in range(num_linhas_template, num_sprints):
                        # Duplica a última linha template
                        nova_linha_idx = duplicar_linha_tabela(table, ultima_linha_template_idx)
                        linhas_template_sprint.append(nova_linha_idx)
                        print(f"[DEBUG] Tabela {table_idx}: Criada nova linha {nova_linha_idx}")
                
                # Preenche as linhas necessárias
                for sprint_idx in range(num_sprints):
                    sprint_data = dados_sprints[sprint_idx]
                    if sprint_idx < len(linhas_template_sprint):
                        linha_idx = linhas_template_sprint[sprint_idx]
                        linha = table.rows[linha_idx]
                        preencher_linha_com_dados_sprint(linha, sprint_data, tags_sprint)
                        print(f"[DEBUG] Preenchida linha {linha_idx} com dados da sprint {sprint_data.get('sprint', 'N/A')}")
                    else:
                        print(f"[DEBUG] ERRO: Não há linha template suficiente para sprint {sprint_idx}")
                
                # Remove linhas extras que não são necessárias (da última para a primeira)
                # IMPORTANTE: Só remove linhas que NÃO foram preenchidas (linhas extras além das necessárias)
                if len(linhas_template_sprint) > num_sprints:
                    linhas_para_remover = len(linhas_template_sprint) - num_sprints
                    print(f"[DEBUG] Removendo {linhas_para_remover} linha(s) extra(s) da tabela {table_idx}")
                    # Remove da última linha template para a primeira extra
                    # IMPORTANTE: Remove em ordem reversa para não afetar os índices
                    linhas_removidas = []
                    for idx in range(len(linhas_template_sprint) - 1, num_sprints - 1, -1):
                        linha_idx_remover = linhas_template_sprint[idx]
                        try:
                            # Verifica se a linha ainda existe antes de remover
                            if linha_idx_remover < len(table.rows):
                                table._element.remove(table.rows[linha_idx_remover]._element)
                                linhas_removidas.append(linha_idx_remover)
                                print(f"[DEBUG] Removida linha {linha_idx_remover} da tabela {table_idx}")
                        except Exception as e:
                            print(f"[DEBUG] Erro ao remover linha {linha_idx_remover}: {e}")
                    print(f"[DEBUG] Total de linhas removidas: {len(linhas_removidas)}")
            else:
                print(f"[DEBUG] Tabela {table_idx}: Nenhuma linha com tags de sprint encontrada")
        
    
    # -----------------------------
    # 2) PROFISSIONAIS / ITEM 7
    # -----------------------------
    # Processa profissionais em tabelas usando tags numeradas
    tags_prof = {
        '{PROF_TIPO}': 'tipo',
        '{PROF_QTD}': 'quantidade',
        '{PROF_QUANTIDADE}': 'quantidade',
        '{PROF_HORAS}': 'horas',
        '{PORCENTAGEM}': 'percentual',
    }
    
    if dados_sprints:
        print(f"[DEBUG] Processando profissionais...")
        
        for table_idx, table in enumerate(doc.tables):
            # Verifica se a tabela usa tags numeradas (ex: {SPRINT_ID_1}, {PROF_TIPO_1_1})
            usa_tags_numeradas = False
            for row in table.rows:
                sprint_num = identificar_sprint_num_na_linha(row)
                if sprint_num is not None:
                    usa_tags_numeradas = True
                    break
            
            if not usa_tags_numeradas:
                # Processa com tags genéricas (compatibilidade com modelos antigos)
                linhas_template_prof = []
                for row_idx, row in enumerate(table.rows):
                    if linha_contem_tag_profissional(row):
                        linhas_template_prof.append((row_idx, row))
                
                if not linhas_template_prof:
                    continue
                
                print(f"[DEBUG] Tabela {table_idx}: Processando com tags genéricas")
                # ... (código antigo para compatibilidade)
                continue
            
            # Processa usando tags numeradas
            print(f"[DEBUG] Tabela {table_idx}: Processando com tags numeradas")
            
            # Agrupa linhas por sprint usando tags numeradas
            grupos_sprint = {}  # {sprint_num: [(row_idx, row, prof_num), ...]}
            
            for row_idx, row in enumerate(table.rows):
                # Pula o cabeçalho (linha 0)
                if row_idx == 0:
                    continue
                # Verifica tags numeradas primeiro
                sprint_num = identificar_sprint_num_na_linha(row)
                if sprint_num is not None:
                    prof_num = identificar_prof_num_na_linha(row, sprint_num)
                    if prof_num is None:
                        prof_num = 1  # Default se não encontrar
                    
                    if sprint_num not in grupos_sprint:
                        grupos_sprint[sprint_num] = []
                    grupos_sprint[sprint_num].append((row_idx, row, prof_num))
                    print(f"[DEBUG] Tabela {table_idx}: Linha {row_idx} tem tags numeradas - Sprint {sprint_num}, Prof {prof_num}")
            
            # Se não encontrou linhas com tags numeradas, tenta identificar linhas COMPLETAMENTE VAZIAS
            # (útil quando o template tem linhas vazias sem tags, mas NÃO preenche linhas com dados normais)
            if not grupos_sprint and len(table.rows) > 1 and dados_sprints:
                print(f"[DEBUG] Tabela {table_idx}: Nenhuma tag numerada encontrada, procurando linhas completamente vazias")
                # Pula o cabeçalho (primeira linha) e procura linhas COMPLETAMENTE VAZIAS
                linha_atual = 1
                for sprint_idx, sprint_data in enumerate(dados_sprints):
                    sprint_num = sprint_idx + 1
                    if linha_atual < len(table.rows):
                        row = table.rows[linha_atual]
                        # Verifica se TODAS as células estão vazias (linha template)
                        todas_vazias = True
                        for cell in row.cells:
                            texto_cell = cell.text.strip()
                            # Se a célula tem conteúdo significativo (mais de 2 caracteres), não é template
                            if len(texto_cell) > 2:
                                todas_vazias = False
                                break
                        
                        if todas_vazias:
                            print(f"[DEBUG] Tabela {table_idx}: Linha {linha_atual} está completamente vazia, adicionando tags numeradas")
                            # Adiciona tags temporárias para que o preenchimento funcione
                            if len(row.cells) > 0:
                                # Limpa e adiciona tag na primeira célula
                                row.cells[0].paragraphs[0].clear()
                                row.cells[0].paragraphs[0].add_run(f'{{SPRINT_ID_{sprint_num}}}')
                            if len(row.cells) > 1:
                                row.cells[1].paragraphs[0].clear()
                                row.cells[1].paragraphs[0].add_run(f'{{SPRINT_TIPO_{sprint_num}}}')
                            if len(row.cells) > 2:
                                row.cells[2].paragraphs[0].clear()
                                row.cells[2].paragraphs[0].add_run(f'{{PROF_TIPO_{sprint_num}_1}}')
                            if len(row.cells) > 3:
                                row.cells[3].paragraphs[0].clear()
                                row.cells[3].paragraphs[0].add_run(f'{{PROF_QTD_{sprint_num}_1}}')
                            if len(row.cells) > 4:
                                row.cells[4].paragraphs[0].clear()
                                row.cells[4].paragraphs[0].add_run(f'{{PROF_HORAS_{sprint_num}_1}}')
                            
                            if sprint_num not in grupos_sprint:
                                grupos_sprint[sprint_num] = []
                            grupos_sprint[sprint_num].append((linha_atual, row, 1))
                            linha_atual += 1
                        else:
                            # Linha tem dados, pula para próxima
                            linha_atual += 1
                    else:
                        break
            
            print(f"[DEBUG] Tabela {table_idx}: Encontrados {len(grupos_sprint)} grupo(s) de sprint")
            
            # Processa cada sprint
            linhas_para_remover = []
            for sprint_idx, sprint_data in enumerate(dados_sprints):
                sprint_num = sprint_idx + 1  # Tags numeradas começam em 1
                sprint_id = str(sprint_data.get('sprint', ''))
                profissionais = dados_profissionais.get(sprint_id, [])
                
                if sprint_num not in grupos_sprint:
                    print(f"[DEBUG] Tabela {table_idx}: Sprint {sprint_num} não encontrada no template")
                    continue
                
                linhas_grupo = grupos_sprint[sprint_num]
                num_profissionais = len(profissionais) if profissionais else 0
                num_linhas_template = len(linhas_grupo)
                
                print(f"[DEBUG] Tabela {table_idx}: Sprint {sprint_num} ({sprint_id}): {num_profissionais} profissional(is), {num_linhas_template} linha(s) template")
                
                # IMPORTANTE: NÃO força preenchimento em células sem tags
                # Só preenche onde há tags explícitas para evitar sobrescrever dados normais
                # O preenchimento já é feito pela função preencher_tags_numeradas_item7
                
                # Preenche linhas existentes
                # IMPORTANTE: Se não houver profissionais, ainda preenche as tags de sprint
                if num_profissionais > 0:
                    # Tem profissionais: preenche normalmente
                    for idx in range(min(num_profissionais, num_linhas_template)):
                        row_idx, row, prof_num_template = linhas_grupo[idx]
                        prof_data = profissionais[idx]
                        prof_num_real = idx + 1  # Profissionais numerados começam em 1
                        primeira_linha_grupo = (idx == 0)  # Primeira linha do grupo de sprint
                        
                        preencher_tags_numeradas_item7(row, sprint_data, prof_data, sprint_num, prof_num_real, primeira_linha_grupo)
                        print(f"[DEBUG] Tabela {table_idx}: Preenchida linha {row_idx} - Sprint {sprint_num}, Profissional {prof_num_real} ({prof_data.get('tipo', 'N/A')})")
                    
                    # Remove linhas extras deste grupo (da última para a primeira)
                    if num_linhas_template > num_profissionais:
                        linhas_remover_grupo = linhas_grupo[num_profissionais:]
                        for row_idx_remover, row_remover, _ in linhas_remover_grupo:
                            linhas_para_remover.append((row_idx_remover, row_remover))
                        print(f"[DEBUG] Tabela {table_idx}: Marcadas {len(linhas_remover_grupo)} linha(s) para remoção do grupo da sprint {sprint_num}")
                else:
                    # Não tem profissionais: preenche apenas tags de sprint na primeira linha
                    if linhas_grupo:
                        row_idx, row, _ = linhas_grupo[0]
                        primeira_linha_grupo = True
                        # Preenche apenas tags de sprint, sem dados de profissional
                        preencher_tags_numeradas_item7(row, sprint_data, None, sprint_num, 1, primeira_linha_grupo)
                        print(f"[DEBUG] Tabela {table_idx}: Preenchida linha {row_idx} - Sprint {sprint_num} (sem profissionais, apenas tags de sprint)")
                        
                        # Remove linhas extras deste grupo (mantém apenas a primeira)
                        if num_linhas_template > 1:
                            linhas_remover_grupo = linhas_grupo[1:]
                            for row_idx_remover, row_remover, _ in linhas_remover_grupo:
                                linhas_para_remover.append((row_idx_remover, row_remover))
                            print(f"[DEBUG] Tabela {table_idx}: Marcadas {len(linhas_remover_grupo)} linha(s) para remoção do grupo da sprint {sprint_num} (sem profissionais)")
                
                # Remove linhas extras deste grupo (da última para a primeira)
                if num_linhas_template > num_profissionais:
                    linhas_remover_grupo = linhas_grupo[num_profissionais:]
                    for row_idx, row, _ in linhas_remover_grupo:
                        linhas_para_remover.append((row_idx, row))
                    print(f"[DEBUG] Tabela {table_idx}: Marcadas {len(linhas_remover_grupo)} linha(s) para remoção do grupo da sprint {sprint_num}")
            
            # Remove linhas extras (da última para a primeira para evitar problemas de índice)
            linhas_para_remover.sort(key=lambda x: x[0], reverse=True)
            for row_idx, row in linhas_para_remover:
                try:
                    table._element.remove(row._element)
                    print(f"[DEBUG] Tabela {table_idx}: Removida linha {row_idx}")
                except Exception as e:
                    print(f"[DEBUG] Erro ao remover linha {row_idx}: {e}")
            
            # Remove grupos de sprint que não existem nos dados
            for sprint_num in grupos_sprint.keys():
                if sprint_num > len(dados_sprints):
                    linhas_grupo = grupos_sprint[sprint_num]
                    for row_idx, row, _ in linhas_grupo:
                        try:
                            table._element.remove(row._element)
                            print(f"[DEBUG] Tabela {table_idx}: Removida linha {row_idx} de sprint {sprint_num} inexistente")
                        except Exception as e:
                            print(f"[DEBUG] Erro ao remover linha {row_idx}: {e}")

        # Compatibilidade: substitui tags fora de tabelas com o primeiro profissional encontrado
        primeiro_prof = None
        for sprint_data in dados_sprints:
            sprint_id = str(sprint_data.get('sprint', ''))
            profissionais = dados_profissionais.get(sprint_id, [])
            if profissionais:
                primeiro_prof = profissionais[0]
                break
        
        if primeiro_prof:
            for tag, campo in tags_prof.items():
                if campo == 'percentual':
                    valor = (
                        primeiro_prof.get('percentual') or
                        primeiro_prof.get('alocacao') or
                        primeiro_prof.get('porcentagem') or
                        ''
                    )
                else:
                    valor = primeiro_prof.get(campo, '')
                substituir_texto_em_documento(doc, tag, str(valor))
        
        # Substitui tags de sprint restantes apenas em parágrafos (fora das tabelas)
        primeira_sprint = dados_sprints[0]
        for tag, campo in tags_sprint.items():
            valor = str(primeira_sprint.get(campo, ''))
            for paragraph in doc.paragraphs:
                substituir_texto_em_paragrafo(paragraph, tag, valor)
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    substituir_texto_em_paragrafo(paragraph, tag, valor)
                for paragraph in section.footer.paragraphs:
                    substituir_texto_em_paragrafo(paragraph, tag, valor)

        # -----------------------------
        # 3) TIPO DA DEMANDA - CHECKBOX
        # -----------------------------
        # Alguns modelos possuem, no item "Tipo da Demanda", caixinhas (☐, U+2610)
        # para: Descoberta, Construção, Design, Manutenção, Arquitetura, Monitoramento.
        # A tabela é fixa, então usamos as posições conhecidas das células de check.
        try:
            # Coleta todos os tipos distintos das sprints
            tipos_sprints = set()
            for s in dados_sprints:
                tipo_val = str(s.get('tipo', '') or '').strip()
                if tipo_val:
                    tipos_sprints.add(tipo_val)

            if tipos_sprints:
                print(f"[DEBUG] Marcando tipo(s) da demanda com base nas sprints: {tipos_sprints}")

                # Normaliza acentuação para comparação
                def normalizar(t: str) -> str:
                    return (
                        t.replace('ã', 'a')
                         .replace('á', 'a')
                         .replace('â', 'a')
                         .replace('é', 'e')
                         .replace('ê', 'e')
                         .replace('í', 'i')
                         .replace('ó', 'o')
                         .replace('ô', 'o')
                         .replace('ú', 'u')
                         .replace('ç', 'c')
                    ).lower()

                # Map de categorias que podem aparecer nas sprints -> chave usada no template
                # (aqui usamos apenas palavras-chave para bater com o texto vindo do Redmine)
                mapa_categoria = {
                    'descoberta': 'descoberta',
                    'construcao': 'construcao',
                    'construcao.': 'construcao',
                    'design': 'design',
                    'manutencao': 'manutencao',
                    'arquitetura': 'arquitetura',
                    'monitoramento': 'monitoramento',
                }

                tipos_selecionados = set()
                for t in tipos_sprints:
                    t_norm = normalizar(t)
                    for chave, categoria in mapa_categoria.items():
                        if chave in t_norm:
                            tipos_selecionados.add(categoria)

                print(f"[DEBUG] Categorias de tipo selecionadas: {tipos_selecionados}")

                if tipos_selecionados:
                    # Localiza a tabela que contém "Tipo da Demanda"
                    tabela_tipo = None
                    for table in doc.tables:
                        encontrou = False
                        for row in table.rows:
                            for cell in row.cells:
                                if 'Tipo da Demanda' in (cell.text or ''):
                                    tabela_tipo = table
                                    encontrou = True
                                    break
                            if encontrou:
                                break
                        if tabela_tipo:
                            break

                    if tabela_tipo:
                        # Pela inspeção do modelo:
                        # linha 1: [0]=☐ Descoberta, [2]=☐ Design, [4]=☐ Arquitetura
                        # linha 2: [0]=☐ Construção, [2]=☐ Manutenção, [4]=☐ Monitoramento
                        checks = []
                        rows = list(tabela_tipo.rows)
                        if len(rows) >= 3:
                            # Cabeçalho está na linha 0; as opções estão nas linhas 1 e 2
                            if len(rows[1].cells) >= 6 and len(rows[2].cells) >= 6:
                                checks = [
                                    # (linha, coluna, categoria)
                                    (1, 0, 'descoberta'),
                                    (2, 0, 'construcao'),
                                    (1, 2, 'design'),
                                    (2, 2, 'manutencao'),
                                    (1, 4, 'arquitetura'),
                                    (2, 4, 'monitoramento'),
                                ]

                        # Primeiro, reseta todos os checkboxes para "não marcado" (☐)
                        for (r, c, _cat) in checks:
                            try:
                                cell = rows[r].cells[c]
                                for p in cell.paragraphs:
                                    substituir_texto_em_paragrafo(p, '☒', '☐')
                                    substituir_texto_em_paragrafo(p, '', '☐')  # fallback
                            except Exception:
                                continue

                        # Depois, marca apenas os tipos que aparecem nas sprints
                        for (r, c, categoria) in checks:
                            if categoria in tipos_selecionados:
                                try:
                                    cell = rows[r].cells[c]
                                    for p in cell.paragraphs:
                                        substituir_texto_em_paragrafo(p, '☐', '☒')
                                        substituir_texto_em_paragrafo(p, '', '☒')  # fallback
                                except Exception:
                                    continue
        except Exception as e:
            print(f"[DEBUG] [WARN] Erro ao marcar tipo da demanda: {e}")

    return doc


def salvar_documento(doc: Document, caminho_saida: str) -> str:
    """
    Salva o documento gerado.
    
    Args:
        doc: Documento Word
        caminho_saida: Caminho onde salvar o arquivo
        
    Returns:
        Caminho do arquivo salvo
    """
    doc.save(caminho_saida)
    return caminho_saida

